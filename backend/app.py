from __future__ import annotations

import json
import io
import base64
import hashlib
import os
import platform
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import threading
import time
import urllib.error
import urllib.request
import unicodedata
import wave
import zipfile
from datetime import datetime
from pathlib import Path

if os.getenv("TRACE_DATA_DIR"):
    APP_DIR = Path(os.environ["TRACE_DATA_DIR"])
elif platform.system() == "Windows":
    APP_DIR = Path(os.getenv("LOCALAPPDATA", Path.home())) / "TRACE-AI"
else:
    APP_DIR = Path.cwd() / ".trace-data"

APP_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = APP_DIR / "trace_memory.db"
WEB_DIR = Path(__file__).resolve().parent.parent / "dist"
OLLAMA_URL = "http://127.0.0.1:11434"
DEFAULT_MODEL = os.getenv("TRACE_MODEL", "qwen3.5:2b-q4_K_M")
VOICE_DIR = APP_DIR / "voice"
COMMAND_VOICE_MODEL = VOICE_DIR / "ggml-large-v3-turbo-q5_0.bin"
VOICE_RELEASE = "v1.9.1"
VOICE_BINARY_URL = f"https://github.com/ggml-org/whisper.cpp/releases/download/{VOICE_RELEASE}/whisper-bin-x64.zip"
VOICE_BINARY_SHA256 = "7d8be46ecd31828e1eb7a2ecdd0d6b314feafd82163038ab6092594b0a063539"
COMMAND_MODEL_URL = "https://huggingface.co/ggerganov/whisper.cpp/resolve/main/ggml-large-v3-turbo-q5_0.bin"
COMMAND_MODEL_SHA1 = "e050f7970618a659205450ad97eb95a18d69c9ee"
TECHNICAL_PROMPT = "TRACE, E aí Trace, Acorde Trace, Descanse Trace, JavaScript, TypeScript, Python, Java, React, Node.js, VS Code, GitHub, Ollama, software, programação."
TTS_DIR = APP_DIR / "tts"
TTS_RUNTIME = TTS_DIR / "runtime"
TTS_MODEL = TTS_DIR / "pt_BR-faber-medium.onnx"
TTS_CONFIG = TTS_DIR / "pt_BR-faber-medium.onnx.json"
TTS_MODEL_URL = "https://huggingface.co/rhasspy/piper-voices/resolve/main/pt/pt_BR/faber/medium/pt_BR-faber-medium.onnx"
TTS_CONFIG_URL = TTS_MODEL_URL + ".json"
DOCUMENT_RUNTIME = APP_DIR / "documents" / "runtime"
_piper_voice = None
_piper_lock = threading.Lock()
_whisper_gate = threading.Lock()
_whisper_process_lock = threading.Lock()
_active_whisper: subprocess.Popen | None = None

SYSTEM_PROMPT = """Você é TRACE, um assistente pessoal local. Fale em português brasileiro.
Sua personalidade é inteligente, calma, levemente irônica e amigável, como um amigo confiável.
Seja direto, útil e natural. Prefira respostas curtas, a menos que o usuário peça detalhes.
Como assistente por voz, comece pela informação ou ação principal. Evite listas longas e introduções.
Em conversa comum, responda em no máximo três frases, salvo quando o usuário pedir detalhes.
Não repita sua identidade, não se reapresente e não transforme cumprimentos simples em discursos.
Use o contexto recente para entender referências como “isso”, “de novo” e “continue”.
Você tem humor discreto e personalidade própria, mas nunca força ironias.
Quando o usuário pedir ajuda com programação, aja como um engenheiro de software experiente: explique o erro, proponha código completo e preserve o contexto dos arquivos anexados.
Em código, priorize soluções executáveis, seguras e fáceis de testar. Não omita partes essenciais com reticências.
Quando o usuário pedir para corrigir, reescrever ou editar um documento anexado, devolva o conteúdo final completo, pronto para ser salvo, sem dizer que enviará um arquivo depois.
Nunca diga que é JARVIS nem imite personagens protegidos.
Não afirme ter executado uma ação se nenhuma ferramenta confirmou a execução.
Não exponha raciocínio interno. Responda somente com a resposta final.
Não use emojis, emoticons, ações entre parênteses ou descrições de expressão/gestos.
Não use Markdown em respostas faladas. Nunca descreva um emoji em palavras.
"""

SAFE_APPS = {
    "bloco de notas": "notepad.exe", "notepad": "notepad.exe",
    "calculadora": "calc.exe", "explorador": "explorer.exe",
    "explorador de arquivos": "explorer.exe", "paint": "mspaint.exe",
    "prompt de comando": "cmd.exe", "terminal": "cmd.exe",
}


def db() -> sqlite3.Connection:
    connection = sqlite3.connect(DB_PATH)
    connection.execute("CREATE TABLE IF NOT EXISTS conversation (id INTEGER PRIMARY KEY, role TEXT NOT NULL, content TEXT NOT NULL, created_at DATETIME DEFAULT CURRENT_TIMESTAMP)")
    return connection


def memory_count() -> int:
    with db() as connection:
        return connection.execute("SELECT COUNT(*) FROM conversation").fetchone()[0]


def is_sound_annotation(text: str) -> bool:
    return bool(re.fullmatch(r"[\[(]\s*(?:risos?|música|aplausos?|silêncio|som[^\])]*|inaudível|ruído)\s*[\])]?[.!]?", text.strip(), flags=re.I))


def recent_messages(limit: int = 10) -> list[dict[str, str]]:
    with db() as connection:
        rows = connection.execute("SELECT role, content FROM conversation ORDER BY id DESC LIMIT ?", (limit * 2,)).fetchall()
    filtered, skip_reply = [], False
    for role, content in reversed(rows):
        if role == "user" and is_sound_annotation(content):
            skip_reply = True
            continue
        if role == "assistant" and skip_reply:
            skip_reply = False
            continue
        skip_reply = False
        filtered.append({"role": role, "content": content})
    return filtered[-limit:]


def conversation_history(limit: int = 40) -> list[dict[str, str]]:
    return recent_messages(max(1, min(limit, 100)))


def remember(role: str, content: str) -> None:
    with db() as connection:
        connection.execute("INSERT INTO conversation(role, content) VALUES (?, ?)", (role, content))


def find_safe_action(text: str) -> tuple[str, str] | None:
    lowered = text.lower().strip()
    if not re.search(r"\b(abra|abrir|inicie|iniciar)\b", lowered):
        return None
    for name, executable in SAFE_APPS.items():
        if name in lowered:
            return name, executable
    return None


def run_safe_action(action: tuple[str, str]) -> str:
    name, executable = action
    if platform.system() != "Windows":
        return f"A abertura de {name} está preparada para o Windows e foi simulada neste ambiente."
    subprocess.Popen([executable], shell=False)
    return f"{name.capitalize()} aberto."


def request_json(url: str, payload: dict | None = None, timeout: float = 2) -> dict:
    data = json.dumps(payload).encode("utf-8") if payload is not None else None
    request = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(request, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def ollama_available() -> bool:
    try:
        request_json(f"{OLLAMA_URL}/api/tags", timeout=1.5)
        return True
    except (OSError, urllib.error.URLError, json.JSONDecodeError):
        return False


def installed_models() -> set[str]:
    try:
        result = request_json(f"{OLLAMA_URL}/api/tags", timeout=2)
        return {str(item.get("name") or item.get("model") or "").removesuffix(":latest") for item in result.get("models", [])}
    except (OSError, urllib.error.URLError, json.JSONDecodeError):
        return set()


def model_installed(model: str = DEFAULT_MODEL) -> bool:
    return model.removesuffix(":latest") in installed_models()


def install_ollama_runtime() -> bool:
    if ollama_available():
        return True
    if platform.system() != "Windows" or shutil.which("winget") is None:
        return False
    creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    commands = [
        ["winget", "install", "--id", "Ollama.Ollama", "-e", "--scope", "user", "--silent", "--accept-package-agreements", "--accept-source-agreements"],
        ["winget", "install", "--id", "Ollama.Ollama", "-e", "--silent", "--accept-package-agreements", "--accept-source-agreements"],
    ]
    try:
        for command in commands:
            result = subprocess.run(command, capture_output=True, timeout=900, creationflags=creation_flags)
            if result.returncode == 0:
                break
        candidates = [
            Path(os.getenv("LOCALAPPDATA", "")) / "Programs" / "Ollama" / "ollama.exe",
            Path(os.getenv("LOCALAPPDATA", "")) / "Ollama" / "ollama.exe",
        ]
        executable = shutil.which("ollama") or next((str(path) for path in candidates if path.is_file()), None)
        if executable:
            subprocess.Popen([executable, "serve"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, creationflags=creation_flags)
        for _ in range(40):
            if ollama_available():
                return True
            time.sleep(.5)
    except (OSError, subprocess.TimeoutExpired):
        return False
    return ollama_available()


def install_model(model: str = DEFAULT_MODEL) -> tuple[bool, str]:
    if not ollama_available() and not install_ollama_runtime():
        return False, "Não foi possível preparar o motor local. Verifique o Instalador de Aplicativo do Windows."
    try:
        request_json(f"{OLLAMA_URL}/api/pull", {"model": model, "stream": False}, timeout=1800)
        return model_installed(model), "Modelo local instalado e pronto."
    except (OSError, urllib.error.URLError, json.JSONDecodeError):
        return False, "Não foi possível baixar o modelo. Verifique sua internet e tente novamente."


def warm_model(model: str = DEFAULT_MODEL) -> None:
    if not model_installed(model):
        return
    try:
        request_json(f"{OLLAMA_URL}/api/generate", {"model": model, "prompt": "", "keep_alive": "60s", "stream": False}, timeout=120)
    except (OSError, urllib.error.URLError, json.JSONDecodeError):
        return


def clean_reply(text: str) -> str:
    text = "".join(char for char in text if unicodedata.category(char) != "So" and char not in "\ufe0f\u200d")
    text = re.sub(r"\*{1,3}|_{1,3}|`{1,3}", "", text)
    text = re.sub(r"\((?:sorrindo|rindo|pensando|pausa|suspiro)[^)]*\)", "", text, flags=re.I)
    text = re.sub(r"\[(?:risos?|música|som[^]]*|pausa|suspiro)[^]]*\]", "", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip()
    return re.sub(r"\s+([,.;:!?])", r"\1", text)


def personalized_prompt(personalization: dict | None) -> str:
    data = personalization or {}
    name = re.sub(r"[^\wÀ-ÿ -]", "", str(data.get("name", "")))[:40].strip()
    style = str(data.get("style", "direct"))
    style_instruction = {
        "friendly": "Use um tom mais próximo e conversador, sem enrolar.",
        "technical": "Seja técnico e detalhado quando isso ajudar, explicando termos importantes.",
        "direct": "Seja direto, natural e objetivo.",
    }.get(style, "Seja direto, natural e objetivo.")
    identity = f"O nome preferido do usuário é {name}. Use-o com moderação." if name else "Não presuma o nome do usuário."
    return f"{SYSTEM_PROMPT}\n{identity}\n{style_instruction}"


def suggested_actions(message: str, attachment_names: list[str], reply: str) -> list[str]:
    lowered = message.casefold()
    if attachment_names:
        return ["Resumir", "Salvar em PDF", "Copiar resposta"]
    if re.search(r"\b(código|programa(?:ção|r)|javascript|typescript|python|java|react|erro|bug)\b", lowered):
        return ["Propor correção", "Aprofundar", "Copiar resposta"]
    if re.search(r"\b(abra|inicie|ative|rotina|modo)\b", lowered):
        return ["Abrir interface", "Copiar resposta"]
    if len(reply) > 420:
        return ["Resumir", "Aprofundar", "Copiar resposta"]
    return ["Aprofundar", "Salvar em PDF", "Copiar resposta"]


def validate_wav(audio: bytes) -> tuple[bool, str]:
    if len(audio) < 44 or len(audio) > 16 * 1024 * 1024:
        return False, "A gravação está vazia ou é muito grande."
    try:
        with wave.open(io.BytesIO(audio), "rb") as wav:
            if wav.getnchannels() != 1 or wav.getsampwidth() != 2:
                return False, "O áudio precisa estar em WAV mono de 16 bits."
            if wav.getnframes() < wav.getframerate() // 3:
                return False, "A gravação ficou curta demais."
    except (wave.Error, EOFError):
        return False, "O formato da gravação não foi reconhecido."
    return True, ""


def voice_binary() -> Path | None:
    if not VOICE_DIR.exists():
        return None
    return next(VOICE_DIR.rglob("whisper-cli.exe"), None)


def voice_ready() -> bool:
    return voice_binary() is not None and COMMAND_VOICE_MODEL.is_file()


def stop_active_whisper() -> bool:
    """Interrompe somente a transcricao iniciada por esta instancia do TRACE."""
    with _whisper_process_lock:
        process = _active_whisper
    if process is None or process.poll() is not None:
        return False
    try:
        process.terminate()
        process.wait(timeout=1.5)
    except (OSError, subprocess.TimeoutExpired):
        try:
            process.kill()
        except OSError:
            pass
    return True


def cleanup_orphaned_whisper() -> None:
    """Remove processos antigos do Whisper cujo comando aponta para os dados do TRACE."""
    if platform.system() != "Windows":
        return
    command = (
        "Get-CimInstance Win32_Process -Filter \"Name = 'whisper-cli.exe'\" | "
        "Where-Object { $_.CommandLine -like '*TRACE-AI*voice*' } | "
        "ForEach-Object { Stop-Process -Id $_.ProcessId -Force -ErrorAction SilentlyContinue }"
    )
    try:
        subprocess.run(
            ["powershell.exe", "-NoProfile", "-NonInteractive", "-Command", command],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=8,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
    except (OSError, subprocess.TimeoutExpired):
        pass


def download_file(url: str, destination: Path) -> None:
    request = urllib.request.Request(url, headers={"User-Agent": "TRACE-AI/0.4"})
    temporary = destination.with_suffix(destination.suffix + ".download")
    with urllib.request.urlopen(request, timeout=1800) as response, temporary.open("wb") as output:
        while chunk := response.read(1024 * 1024):
            output.write(chunk)
    temporary.replace(destination)


def install_voice_engine() -> tuple[bool, str]:
    VOICE_DIR.mkdir(parents=True, exist_ok=True)
    try:
        if voice_binary() is None:
            archive = VOICE_DIR / "whisper-windows.zip"
            download_file(VOICE_BINARY_URL, archive)
            if hashlib.sha256(archive.read_bytes()).hexdigest() != VOICE_BINARY_SHA256:
                archive.unlink(missing_ok=True)
                return False, "O programa de voz baixado não passou na verificação de integridade."
            with zipfile.ZipFile(archive) as package:
                package.extractall(VOICE_DIR / "runtime")
            archive.unlink(missing_ok=True)
        if not COMMAND_VOICE_MODEL.is_file():
            download_file(COMMAND_MODEL_URL, COMMAND_VOICE_MODEL)
        if COMMAND_VOICE_MODEL.stat().st_size < 520_000_000 or hashlib.sha1(COMMAND_VOICE_MODEL.read_bytes()).hexdigest() != COMMAND_MODEL_SHA1:
            COMMAND_VOICE_MODEL.unlink(missing_ok=True)
            return False, "O modelo de transcrição avançada veio incompleto."
        if not voice_ready():
            return False, "O executável do Whisper não foi encontrado após a instalação."
        return True, "Motor de voz local instalado e pronto."
    except (OSError, urllib.error.URLError, zipfile.BadZipFile):
        return False, "Não foi possível baixar o motor de voz. Verifique a internet e tente novamente."


def tts_ready() -> bool:
    return (TTS_RUNTIME / "piper" / "__init__.py").is_file() and TTS_MODEL.is_file() and TTS_CONFIG.is_file()


def install_tts_engine() -> tuple[bool, str]:
    TTS_DIR.mkdir(parents=True, exist_ok=True)
    TTS_RUNTIME.mkdir(parents=True, exist_ok=True)
    try:
        if not (TTS_RUNTIME / "piper" / "__init__.py").is_file():
            creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            result = subprocess.run(
                [sys.executable, "-m", "pip", "install", "--disable-pip-version-check", "--no-warn-script-location", "--target", str(TTS_RUNTIME), "piper-tts==1.6.0"],
                capture_output=True, timeout=900, creationflags=creation_flags,
            )
            if result.returncode != 0:
                return False, "Não foi possível preparar a voz neural local."
        if not TTS_MODEL.is_file():
            download_file(TTS_MODEL_URL, TTS_MODEL)
        if not TTS_CONFIG.is_file():
            download_file(TTS_CONFIG_URL, TTS_CONFIG)
        if not tts_ready():
            return False, "A voz neural foi baixada, mas ficou incompleta."
        return True, "Voz neural brasileira instalada e pronta."
    except (OSError, urllib.error.URLError, subprocess.TimeoutExpired):
        return False, "Não foi possível preparar a voz neural local."


def document_reader_ready() -> bool:
    return (DOCUMENT_RUNTIME / "pypdf" / "__init__.py").is_file() and (DOCUMENT_RUNTIME / "docx" / "__init__.py").is_file()


def install_document_reader() -> tuple[bool, str]:
    if document_reader_ready():
        return True, "Leitor local de PDF, DOCX e OCR visual já está instalado."
    DOCUMENT_RUNTIME.mkdir(parents=True, exist_ok=True)
    try:
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "--disable-pip-version-check", "--no-warn-script-location", "--target", str(DOCUMENT_RUNTIME), "pypdf>=5,<7", "python-docx>=1.1,<2", "pymupdf>=1.24,<2"],
            capture_output=True, timeout=600, creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        if result.returncode != 0 or not document_reader_ready():
            return False, "Não foi possível preparar o leitor local de PDF."
        return True, "Leitor local de PDF, DOCX e OCR visual instalado e pronto."
    except (OSError, subprocess.TimeoutExpired):
        return False, "Não foi possível preparar o leitor local de PDF."


def attachment_context(items: list[dict]) -> tuple[str, list[str], list[str]]:
    texts, images, names = [], [], []
    for item in items[:4]:
        name = Path(str(item.get("name", "arquivo"))).name[:180]
        kind = str(item.get("type", "")).lower()
        try:
            raw = base64.b64decode(str(item.get("data", "")), validate=True)
        except (ValueError, TypeError):
            continue
        if not raw or len(raw) > 12 * 1024 * 1024:
            continue
        names.append(name)
        if kind in {"png", "jpg", "jpeg", "webp"}:
            images.append(base64.b64encode(raw).decode("ascii"))
        elif kind == "pdf":
            if not document_reader_ready():
                texts.append(f"[PDF {name}: leitor de PDF ainda não autorizado]")
                continue
            runtime = str(DOCUMENT_RUNTIME)
            if runtime not in sys.path:
                sys.path.insert(0, runtime)
            try:
                from pypdf import PdfReader
                content = "\n".join((page.extract_text() or "") for page in PdfReader(io.BytesIO(raw)).pages[:40])
                texts.append(f"[PDF: {name}]\n{content[:50000]}")
                if len(content.strip()) < 120:
                    try:
                        import fitz
                        document = fitz.open(stream=raw, filetype="pdf")
                        for page in list(document)[:6]:
                            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.45, 1.45), alpha=False)
                            images.append(base64.b64encode(pixmap.tobytes("jpeg", jpg_quality=76)).decode("ascii"))
                        texts.append(f"[OCR VISUAL: {name} foi enviado como imagem por não conter texto selecionável]")
                    except (ImportError, RuntimeError, ValueError):
                        texts.append(f"[OCR VISUAL indisponível para {name}]")
            except (OSError, ValueError):
                texts.append(f"[PDF {name}: não foi possível extrair o texto]")
        elif kind == "docx":
            if not document_reader_ready():
                texts.append(f"[DOCX {name}: leitor de documentos ainda não autorizado]")
                continue
            runtime = str(DOCUMENT_RUNTIME)
            if runtime not in sys.path:
                sys.path.insert(0, runtime)
            try:
                from docx import Document
                document = Document(io.BytesIO(raw))
                paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
                tables = [" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows]
                texts.append(f"[DOCX: {name}]\n" + "\n".join([*paragraphs, *tables])[:50000])
            except (ImportError, OSError, ValueError):
                texts.append(f"[DOCX {name}: não foi possível extrair o conteúdo]")
        else:
            content = raw.decode("utf-8", errors="replace")
            texts.append(f"[ARQUIVO: {name}]\n{content[:50000]}")
    return "\n\n".join(texts)[:70000], images, names


def create_docx(text: str) -> bytes | None:
    if not document_reader_ready():
        return None
    try:
        runtime = str(DOCUMENT_RUNTIME)
        if runtime not in sys.path:
            sys.path.insert(0, runtime)
        from docx import Document
        document = Document()
        document.add_heading("Documento revisado pelo TRACE", level=1)
        for block in re.split(r"\n{2,}", text.strip()[:100000]):
            document.add_paragraph(block)
        output = io.BytesIO()
        document.save(output)
        return output.getvalue()
    except (ImportError, OSError, ValueError):
        return None


def synthesize_speech(text: str) -> bytes | None:
    global _piper_voice
    if not tts_ready() or not text.strip():
        return None
    try:
        runtime = str(TTS_RUNTIME)
        if runtime not in sys.path:
            sys.path.insert(0, runtime)
        from piper import PiperVoice, SynthesisConfig
        with _piper_lock:
            if _piper_voice is None:
                _piper_voice = PiperVoice.load(str(TTS_MODEL))
            output = io.BytesIO()
            with wave.open(output, "wb") as wav_file:
                config = SynthesisConfig(volume=0.94, length_scale=0.94, noise_scale=0.56, noise_w_scale=0.72, normalize_audio=True)
                _piper_voice.synthesize_wav(text[:700], wav_file, syn_config=config)
            return output.getvalue()
    except (OSError, ImportError, RuntimeError, TypeError, ValueError):
        return None


def transcribe_whisper(audio: bytes, purpose: str = "command") -> tuple[bool, str]:
    global _active_whisper
    valid, error = validate_wav(audio)
    if not valid:
        return False, error
    executable = voice_binary()
    is_wake = purpose == "wake"
    model = COMMAND_VOICE_MODEL
    if executable is None or not model.is_file():
        return False, "O motor de voz local ainda não foi instalado. Abra as configurações do TRACE para preparar a voz."
    if not _whisper_gate.acquire(blocking=False):
        return False, "O reconhecimento de voz já está processando outra fala."
    try:
        with tempfile.TemporaryDirectory(prefix="trace-voice-") as temporary:
            temp_dir = Path(temporary)
            audio_path = temp_dir / "command.wav"
            output_base = temp_dir / "transcript"
            audio_path.write_bytes(audio)
            creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            transcript_path = output_base.with_suffix(".txt")
            threads = "2" if is_wake else str(max(2, min(4, (os.cpu_count() or 4) // 2)))
            common = [str(executable), "-m", str(model), "-f", str(audio_path), "-l", "pt", "-otxt", "-of", str(output_base), "-nt", "-t", threads]
            results = []
            attempts = [[*common, "-p", TECHNICAL_PROMPT]] if is_wake else [[*common, "-p", TECHNICAL_PROMPT], common]
            for arguments in attempts:
                transcript_path.unlink(missing_ok=True)
                process = subprocess.Popen(arguments, stdout=subprocess.PIPE, stderr=subprocess.PIPE, creationflags=creation_flags)
                with _whisper_process_lock:
                    _active_whisper = process
                try:
                    stdout, stderr = process.communicate(timeout=25 if is_wake else 90)
                except subprocess.TimeoutExpired:
                    process.kill()
                    stdout, stderr = process.communicate()
                    raise
                finally:
                    with _whisper_process_lock:
                        if _active_whisper is process:
                            _active_whisper = None
                results.append((process.returncode, stdout, stderr))
                if transcript_path.is_file() and transcript_path.stat().st_size:
                    break
            text = transcript_path.read_text(encoding="utf-8", errors="replace").strip() if transcript_path.is_file() else ""
            if not text:
                output = results[-1][1].decode("utf-8", errors="replace") if results else ""
                lines = [re.sub(r"^\[[^]]+\]\s*", "", line).strip() for line in output.splitlines() if line.strip()]
                text = " ".join(line for line in lines if not line.startswith(("whisper_", "system_info", "main:"))).strip()
            if not text:
                return False, "Não consegui entender. Fale novamente depois do sinal."
            text = re.sub(r"\s+", " ", text).strip()
            if is_sound_annotation(text) or len(re.sub(r"\W", "", text)) < 2:
                return False, "Não detectei uma frase."
            return True, text[:1000]
    except (OSError, subprocess.TimeoutExpired):
        return False, "O motor de voz local não respondeu."
    finally:
        with _whisper_process_lock:
            _active_whisper = None
        _whisper_gate.release()


def direct_answer(message: str) -> str | None:
    normalized = message.casefold().strip()
    now = datetime.now().astimezone()
    if re.search(r"\b(que horas|qual(?: é| e)? o horário|horas agora|me diga as horas)\b", normalized):
        return f"Agora são {now.strftime('%H:%M')}."
    if re.search(r"\b(que dia é hoje|qual(?: é| e)? a data|data de hoje)\b", normalized):
        return f"Hoje é {now.strftime('%d/%m/%Y')}."
    return None


def chat(message: str, model: str = DEFAULT_MODEL, attachments: list[dict] | None = None, permissions: dict | None = None, profile: str = "balanced", personalization: dict | None = None) -> dict:
    message = message.strip()[:4000]
    permissions = permissions or {}
    context, images, attachment_names = attachment_context(attachments or []) if permissions.get("files") or permissions.get("screen") else ("", [], [])
    if not message and not attachment_names:
        return {"reply": "Não recebi nenhum comando.", "action": None}
    if attachment_names and not message:
        message = "Analise os arquivos anexados e me diga os pontos mais importantes."
    remember("user", message)
    exact_reply = direct_answer(message)
    action = find_safe_action(message)
    action_result = run_safe_action(action) if action and permissions.get("apps") else None
    if exact_reply:
        reply = exact_reply
    elif action and not permissions.get("apps"):
        reply = "O acesso a aplicativos está desativado. Você pode autorizá-lo nas configurações."
    elif action_result:
        reply = action_result
    elif model_installed(model):
        coding = profile == "coding" or bool(re.search(r"\b(código|programa(?:ção|r)|javascript|typescript|python|java|react|erro|bug)\b", message, re.I))
        payload = {
            "model": model[:100], "stream": False, "think": False, "keep_alive": "5m",
            "options": {"num_ctx": 4096 if coding else 2048, "num_predict": 420 if coding else 160, "temperature": 0.28 if coding else 0.38, "repeat_penalty": 1.08},
            "messages": [{"role": "system", "content": personalized_prompt(personalization)}, *recent_messages(8)],
        }
        if context or images:
            payload["messages"][-1] = {"role": "user", "content": f"{message}\n\n{context}".strip(), **({"images": images} if images else {})}
        try:
            reply = clean_reply(request_json(f"{OLLAMA_URL}/api/chat", payload, timeout=120)["message"]["content"])
        except (OSError, urllib.error.URLError, KeyError, json.JSONDecodeError):
            reply = "Meu modelo local demorou a responder. Tente novamente em alguns instantes."
    else:
        reply = "Meu núcleo está funcionando, mas o Qwen 3.5 ainda está sendo preparado. Aguarde a instalação automática."
    remember("assistant", reply)
    return {"reply": reply, "action": action_result, "suggestions": suggested_actions(message, attachment_names, reply)}


def stream_chat(message: str, model: str = DEFAULT_MODEL, attachments: list[dict] | None = None, permissions: dict | None = None, profile: str = "balanced", personalization: dict | None = None):
    message = message.strip()[:4000]
    permissions = permissions or {}
    context, images, attachment_names = attachment_context(attachments or []) if permissions.get("files") or permissions.get("screen") else ("", [], [])
    if not message and not attachment_names:
        yield {"type": "done", "reply": "Não recebi nenhum comando.", "suggestions": []}
        return
    if attachment_names and not message:
        message = "Analise os arquivos anexados e me diga os pontos mais importantes."
    remember("user", message)
    exact_reply = direct_answer(message)
    action = find_safe_action(message)
    action_result = run_safe_action(action) if action and permissions.get("apps") else None
    if exact_reply or action_result or (action and not permissions.get("apps")) or not model_installed(model):
        if exact_reply:
            reply = exact_reply
        elif action_result:
            reply = action_result
        elif action:
            reply = "O acesso a aplicativos está desativado. Você pode autorizá-lo nas configurações."
        else:
            reply = "Meu núcleo está funcionando, mas o Qwen 3.5 ainda está sendo preparado. Aguarde a instalação automática."
        remember("assistant", reply)
        yield {"type": "done", "reply": reply, "action": action_result, "suggestions": suggested_actions(message, attachment_names, reply)}
        return
    coding = profile == "coding" or bool(re.search(r"\b(código|programa(?:ção|r)|javascript|typescript|python|java|react|erro|bug)\b", message, re.I))
    payload = {
        "model": model[:100], "stream": True, "think": False, "keep_alive": "5m",
        "options": {"num_ctx": 4096 if coding else 2048, "num_predict": 420 if coding else 160, "temperature": 0.28 if coding else 0.38, "repeat_penalty": 1.08},
        "messages": [{"role": "system", "content": personalized_prompt(personalization)}, *recent_messages(8)],
    }
    if context or images:
        payload["messages"][-1] = {"role": "user", "content": f"{message}\n\n{context}".strip(), **({"images": images} if images else {})}
    accumulated = ""
    try:
        request = urllib.request.Request(f"{OLLAMA_URL}/api/chat", data=json.dumps(payload).encode("utf-8"), headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(request, timeout=120) as response:
            for raw_line in response:
                if not raw_line.strip():
                    continue
                item = json.loads(raw_line.decode("utf-8"))
                delta = str(item.get("message", {}).get("content", ""))
                if delta:
                    accumulated += delta
                    yield {"type": "delta", "text": delta}
    except (OSError, urllib.error.URLError, json.JSONDecodeError):
        accumulated = "Meu modelo local demorou a responder. Tente novamente em alguns instantes."
    reply = clean_reply(accumulated).strip() or "Não consegui gerar uma resposta completa. Tente novamente."
    remember("assistant", reply)
    yield {"type": "done", "reply": reply, "suggestions": suggested_actions(message, attachment_names, reply)}
